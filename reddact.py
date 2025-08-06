import os
import re
import sys
import json
import uuid
import docx
import fitz 
import spacy
from faker import Faker
from PIL import Image, ImageDraw
import pytesseract
from cryptography.fernet import Fernet
from bs4 import BeautifulSoup, Comment, CData, ProcessingInstruction, Doctype

fake = Faker()
nlp = spacy.load("en_core_web_sm")
regex_patterns = {
    "EMAIL": r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}',
    "SSN": r'\d{3}-\d{2}-\d{4}',
    "IP_ADDRESS": r'\b(?:\d{1,3}\.){3}\d{1,3}\b',
    "CREDIT_CARD": r'\b(?:\d[ -]*?){13,16}\b'
}

def load_key():
    key = os.environ.get('SECRET_ENCRYPTION_KEY')
    if key is None:
        raise ValueError("ERROR: SECRET_ENCRYPTION_KEY is not set in your .env file!")
    return key.encode()

def load_redaction_db():
    if os.path.exists("redaction_db.json"):
        with open("redaction_db.json", 'r', encoding='utf-8') as f:
            return json.load(f)
    return {}

def save_redaction_db(db):
    with open("redaction_db.json", 'w', encoding='utf-8') as f:
        json.dump(db, f, indent=2)


def redact_text(text, redaction_type="blackout", f_cipher=None, db=None):
    audit_log = []
    
    text_to_process = text

    #NER Entities 
    doc = nlp(text_to_process)
    for ent in doc.ents:
        if ent.label_ in ["PERSON", "GPE", "ORG"]:
            original_value = ent.text
            replacement = ""
            if redaction_type == "blackout":
                replacement = "█" * len(original_value)
            elif redaction_type == "synthetic":
                if ent.label_ == "PERSON": replacement = fake.name()
                else: replacement = fake.city()
            elif redaction_type == "reversible":
                if f_cipher is None or db is None: raise ValueError("Cipher/DB missing.")
                redaction_id = f"[REDACT-{uuid.uuid4()}]"
                db[redaction_id] = {"data": f_cipher.encrypt(original_value.encode()).decode(), "type": ent.label_}
                replacement = redaction_id
            
            audit_log.append({"original": original_value, "redacted_to": replacement, "type": ent.label_})
            text = text.replace(original_value, replacement)

    #Regex Patterns
    for pii_type, pattern in regex_patterns.items():
        matches = list(re.finditer(pattern, text))
        for match in matches:
            original_value = match.group(0)
            replacement = ""
            if redaction_type == "blackout":
                replacement = f"[{pii_type}_REDACTED]"
            elif redaction_type == "synthetic":
                if pii_type == "EMAIL": replacement = fake.email()
                else: replacement = f"[{pii_type}_SYNTHETIC]"
            elif redaction_type == "reversible":
                if f_cipher is None or db is None: raise ValueError("Cipher/DB missing.")
                redaction_id = f"[REDACT-{uuid.uuid4()}]"
                db[redaction_id] = {"data": f_cipher.encrypt(original_value.encode()).decode(), "type": pii_type}
                replacement = redaction_id
            
            audit_log.append({"original": original_value, "redacted_to": replacement, "type": pii_type})
            text = text.replace(original_value, replacement)
            
    return text, audit_log

def redact_structured_data(data, audit_collector, redaction_type, f_cipher, db):
    if isinstance(data, dict):
        for key, value in data.items():
            data[key] = redact_structured_data(value, audit_collector, redaction_type, f_cipher, db)
    elif isinstance(data, list):
        for i, item in enumerate(data):
            data[i] = redact_structured_data(item, audit_collector, redaction_type, f_cipher, db)
    elif isinstance(data, str):
        redacted_string, audit_entries = redact_text(data, redaction_type, f_cipher, db)
        if audit_entries:
            audit_collector.extend(audit_entries)
        return redacted_string
    return data


def process_file(file_path, redaction_type="blackout"):
    print(f"Processing '{file_path}' with redaction type: '{redaction_type}'")
    master_audit_log = []
    output_filename = "" 
    original_filename = os.path.basename(file_path)

    f_cipher, redaction_db = None, None
    if redaction_type == "reversible":
        key = load_key()
        f_cipher = Fernet(key)
        redaction_db = load_redaction_db()
        
    try:
        if file_path.lower().endswith(".pdf"):
            doc = fitz.open(file_path)
            for page in doc:
                text, audit = redact_text(page.get_text("text"), redaction_type, f_cipher, redaction_db)
            output_filename = original_filename.replace(".pdf", "_redacted.pdf")

        elif file_path.lower().endswith((".png", ".jpg", ".jpeg")):
            img = Image.open(file_path)
            full_text = pytesseract.image_to_string(img)
            _, audit_entries = redact_text(full_text, redaction_type, f_cipher, redaction_db)
            master_audit_log.extend(audit_entries)

            draw = ImageDraw.Draw(img)
            ocr_data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT)
            for i in range(len(ocr_data['text'])):
                word = ocr_data['text'][i]
                if any(entry['original'] in word or word in entry['original'] for entry in master_audit_log):
                     x, y, w, h = ocr_data['left'][i], ocr_data['top'][i], ocr_data['width'][i], ocr_data['height'][i]
                     draw.rectangle([(x, y), (x + w, y + h)], fill='black')
            ext = original_filename.split('.')[-1]
            output_filename = original_filename.replace(f".{ext}", f"_redacted.{ext}")
            img.save(output_filename)

        elif file_path.lower().endswith(".docx"):
            doc = docx.Document(file_path)
            new_doc = docx.Document()
            for para in doc.paragraphs:
                redacted_para, audit_entries = redact_text(para.text, redaction_type, f_cipher, redaction_db)
                master_audit_log.extend(audit_entries)
                new_doc.add_paragraph(redacted_para)
            output_filename = original_filename.replace(".docx", "_redacted.docx")
            new_doc.save(output_filename)
        
        elif file_path.lower().endswith(".txt"):
            with open(file_path, 'r', encoding='utf-8') as f: text_to_redact = f.read()
            redacted_content, audit_entries = redact_text(text_to_redact, redaction_type, f_cipher, redaction_db)
            master_audit_log.extend(audit_entries)
            output_filename = original_filename.replace(".txt", "_redacted.txt")
            with open(output_filename, 'w', encoding='utf-8') as f: f.write(redacted_content)
            
        elif file_path.endswith(".json"):
            with open(file_path, 'r', encoding='utf-8') as f: data = json.load(f)
            redacted_data = redact_structured_data(data, master_audit_log, redaction_type, f_cipher, redaction_db)
            output_filename = original_filename.replace(".json", "_redacted.json")
            with open(output_filename, 'w', encoding='utf-8') as f: json.dump(redacted_data, f, indent=2)

        elif file_path.endswith(".xml"):
            with open(file_path, 'r', encoding='utf-8') as f: soup = BeautifulSoup(f, 'lxml-xml')
            for tag in soup.find_all(string=True):
                if not str(tag.string).strip() or isinstance(tag, (Comment, CData, ProcessingInstruction, Doctype)): continue
                redacted_string, audit_entries = redact_text(str(tag.string), redaction_type, f_cipher, redaction_db)
                if audit_entries: master_audit_log.extend(audit_entries)
                tag.string.replace_with(redacted_string)
            output_filename = original_filename.replace(".xml", "_redacted.xml")
            with open(output_filename, 'w', encoding='utf-8') as f: f.write(str(soup.prettify()))
        
        else:
            raise ValueError(f"Unsupported file type: {original_filename}")

        print(f"Successfully created redacted file: {output_filename}")

        # -------- Save Databases and Audit Logs ---------
        if redaction_type == "reversible":
            save_redaction_db(redaction_db)
            print("Successfully updated the reversible redaction database.")
        
        if master_audit_log:
            base_name = ".".join(original_filename.split('.')[:-1])
            audit_filename = f"{base_name}_audit.json"
            with open(audit_filename, 'w', encoding='utf-8') as f:
                json.dump(master_audit_log, f, indent=4)
            print(f"Successfully saved audit log to {audit_filename}")
        
        return output_filename

    except Exception as e:
        print(f"ERROR during processing: {e}")
        raise e